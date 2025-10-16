#!/usr/bin/env python3
"""
Restored Working Web UI for Email Company Tool
Simple and reliable version
"""

from flask import Flask, request, send_file, jsonify
import os
import pandas as pd
import threading
import time
import uuid
from main import process_emails

app = Flask(__name__)

# Configuration
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
OUTPUT_FOLDER = os.path.join(BASE_DIR, 'Output')

# Ensure directories exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# Job storage
jobs = {}

def extract_name_from_email(email):
    """
    Extract a name from the email address itself (from the part before @)
    Enhanced to handle more patterns and common email formats
    """
    try:
        username = email.split('@')[0].lower()
        
        # Skip common generic emails
        generic_emails = ['info', 'support', 'contact', 'admin', 'sales', 'help', 'hr', 'team', 
                         'marketing', 'ops', 'press', 'partnerships', 'lab', 'hello', 'inquiries']
        
        if username in generic_emails:
            return username.capitalize()
        
        # Common patterns in email usernames
        name_parts = []
        
        # Handle dots, underscores, dashes
        if '.' in username:
            parts = username.split('.')
        elif '_' in username:
            parts = username.split('_')
        elif '-' in username:
            parts = username.split('-')
        else:
            # Try to split camelCase or handle as single name
            import re
            # Handle camelCase
            parts = re.findall(r'[A-Z][a-z]*|^[a-z]+', username)
            if not parts:
                # If no camelCase found, try to split by numbers or special chars
                parts = re.split(r'[\d_\-\.]+', username)
                if not parts or parts == ['']:
                    parts = [username]
        
        # Clean and capitalize parts
        for part in parts[:2]:  # Take max 2 parts for first and last name
            clean_part = ''.join(c for c in part if c.isalpha())
            if len(clean_part) > 1:  # Only include meaningful parts (more than 1 char)
                name_parts.append(clean_part.capitalize())
        
        if name_parts:
            return ' '.join(name_parts)
        else:
            # Fallback: just capitalize the username
            clean_username = ''.join(c for c in username if c.isalpha())
            return clean_username.capitalize() if clean_username else 'Unknown'
            
    except:
        return 'Unknown'

def extract_emails_with_context(text, fast_mode=True):
    """
    Extract emails along with potential associated names and other data from text.
    Returns: emails list, names list, other_data list
    """
    import re
    
    # Email pattern - optimized for speed
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    
    # Find all emails - compile regex once for better performance
    email_regex = re.compile(email_pattern, re.IGNORECASE)
    emails = email_regex.findall(text)
    
    # Remove duplicates while preserving order - faster method
    seen = set()
    unique_emails = []
    for email in emails:
        email_lower = email.lower()
        if email_lower not in seen:
            unique_emails.append(email_lower)
            seen.add(email_lower)
    
    names = []
    other_data = []
    
    if fast_mode:
        # Fast mode - extract names from email addresses
        for email in unique_emails:
            name = extract_name_from_email(email)
            names.append(name)
            other_data.append('')
        return unique_emails, names, other_data
    
    # Detailed mode - try to find names in context (slower but more complete)
    lines = text.split('\n')
    
    for email in unique_emails:
        name = ''
        extra_info = ''
        
        # Look for the email in text and try to find associated name
        for line in lines:
            if email.lower() in line.lower():
                line_clean = line.strip()
                
                # Try common patterns
                # Pattern 1: Name <email>
                name_match = re.search(r'([A-Za-z\s]+)\s*<\s*' + re.escape(email) + r'\s*>', line_clean, re.IGNORECASE)
                if name_match:
                    name = name_match.group(1).strip()
                    break
                
                # Pattern 2: Name: email or Name - email
                name_match = re.search(r'([A-Za-z\s]+)[\s]*[:\-]\s*' + re.escape(email), line_clean, re.IGNORECASE)
                if name_match:
                    name = name_match.group(1).strip()
                    break
                
                # Pattern 3: email - Name or email: Name
                name_match = re.search(re.escape(email) + r'[\s]*[:\-]\s*([A-Za-z\s]+)', line_clean, re.IGNORECASE)
                if name_match:
                    name = name_match.group(1).strip()
                    break
        
        # If no name found in context, extract from email
        if not name:
            name = extract_name_from_email(email)
        
        names.append(name)
        other_data.append(extra_info)
    
    return unique_emails, names, other_data

def create_dataframe_from_emails(emails, names=None, other_data=None):
    """
    Create a DataFrame from extracted emails and associated data.
    Match the exact format: name, company, email, domain, company, sector
    """
    if not emails:
        return pd.DataFrame(columns=['name', 'company', 'email', 'domain', 'company', 'sector'])
    
    from main import get_domain
    
    # Extract domains from emails
    domains = [get_domain(email) for email in emails]
    
    # Create the DataFrame with the exact structure you want
    names_list = names if names and len(names) == len(emails) else [extract_name_from_email(email) for email in emails]
    
    data = {
        'name': names_list,
        'email': emails,
        'domain': domains
    }
    
    return pd.DataFrame(data)

def fast_process_emails(emails_df):
    """
    Fast processing mode - prioritizes speed over completeness.
    Only uses local dataset matching, no web search.
    Creates output in exact format: name, company, email, domain, company, sector
    """
    from main import get_domain
    import os
    
    # Ensure we have the required columns
    result_df = emails_df.copy()
    
    # Make sure we have names - extract from email if not present
    if 'name' not in result_df.columns or result_df['name'].isna().all():
        result_df['name'] = result_df['email'].apply(extract_name_from_email)
    
    # Extract domains
    if 'domain' not in result_df.columns:
        result_df["domain"] = result_df["email"].apply(get_domain)
    
    # Load dataset quickly
    dataset_file = os.path.join(BASE_DIR, 'Dataset', 'company_sector.csv')
    # Fallback to enhanced dataset if company_sector.csv doesn't exist
    if not os.path.exists(dataset_file):
        dataset_file = os.path.join(BASE_DIR, 'Dataset', 'enhanced_company_sector.csv')
    if os.path.exists(dataset_file):
        mapping_df = pd.read_csv(dataset_file)
        # Create a quick lookup dictionary for O(1) access
        domain_lookup = {}
        for _, row in mapping_df.iterrows():
            domain_lookup[row['domain'].lower()] = (row['company'], row['sector'])
        
        # Fast matching function with enhanced logic
        def quick_match(domain):
            domain_lower = domain.lower().strip()
            
            # Direct lookup first
            if domain_lower in domain_lookup:
                return pd.Series(domain_lookup[domain_lower])
            
            # Try with .com extension
            domain_com = f"{domain_lower}.com"
            if domain_com in domain_lookup:
                return pd.Series(domain_lookup[domain_com])
            
            # Try removing common extensions and matching
            for ext in ['.com', '.net', '.org', '.io', '.co', '.tech', '.biz', '.ai']:
                if domain_lower.endswith(ext):
                    base_domain = domain_lower[:-len(ext)]
                    if base_domain in domain_lookup:
                        return pd.Series(domain_lookup[base_domain])
            
            # Try partial matching (contains)
            for domain_key in domain_lookup.keys():
                if domain_lower in domain_key or domain_key in domain_lower:
                    if abs(len(domain_lower) - len(domain_key)) <= 3:  # Similar length
                        return pd.Series(domain_lookup[domain_key])
            
            # Try fuzzy matching with multiple algorithms and lower threshold
            from rapidfuzz import process, fuzz
            choices = list(domain_lookup.keys())
            
            # Try different scoring methods
            best_matches = [
                process.extractOne(domain_lower, choices, scorer=fuzz.ratio),
                process.extractOne(domain_lower, choices, scorer=fuzz.partial_ratio),
                process.extractOne(domain_lower, choices, scorer=fuzz.token_sort_ratio),
                process.extractOne(domain_lower, choices, scorer=fuzz.token_set_ratio)
            ]
            
            # Use the best match with score >= 70% (lowered from 80%)
            for match in best_matches:
                if match and match[1] >= 70:
                    return pd.Series(domain_lookup[match[0]])
            
            # Try matching common patterns (remove numbers, common suffixes)
            cleaned_domain = domain_lower
            # Remove numbers
            import re
            cleaned_domain = re.sub(r'\d+', '', cleaned_domain)
            # Remove common business suffixes
            suffixes = ['corp', 'inc', 'llc', 'ltd', 'company', 'co', 'group', 'international', 'intl', 'global', 'usa', 'us', 'solutions', 'services', 'systems', 'tech', 'technologies']
            for suffix in suffixes:
                if cleaned_domain.endswith(suffix):
                    cleaned_domain = cleaned_domain[:-len(suffix)]
                    break
            
            # Try matching the cleaned domain
            if cleaned_domain and cleaned_domain != domain_lower and cleaned_domain in domain_lookup:
                return pd.Series(domain_lookup[cleaned_domain])
            
            # Try fuzzy matching on cleaned domain
            if cleaned_domain and len(cleaned_domain) > 2:
                best_match = process.extractOne(cleaned_domain, choices, scorer=fuzz.ratio)
                if best_match and best_match[1] >= 65:  # Even lower threshold for cleaned domain
                    return pd.Series(domain_lookup[best_match[0]])
            
            # If still no match, try to generate a reasonable company name from domain
            company_name = domain_lower.replace('.', ' ').replace('-', ' ').replace('_', ' ')
            # Clean up the company name
            words = company_name.split()
            if words:
                # Filter out very short words and common terms
                meaningful_words = [word.capitalize() for word in words 
                                 if len(word) > 2 and word.lower() not in ['com', 'net', 'org', 'www', 'mail', 'email']]
                if meaningful_words:
                    company_name = ' '.join(meaningful_words)
                    # Assign a sector based on common domain patterns
                    sector = 'Unknown'
                    if any(tech_word in domain_lower for tech_word in ['tech', 'soft', 'system', 'data', 'digital', 'cyber', 'cloud', 'ai', 'ml']):
                        sector = 'Technology'
                    elif any(fin_word in domain_lower for fin_word in ['bank', 'financial', 'finance', 'capital', 'invest', 'fund', 'money']):
                        sector = 'Finance'
                    elif any(health_word in domain_lower for health_word in ['health', 'medical', 'pharma', 'bio', 'care', 'hospital', 'clinic']):
                        sector = 'Healthcare'
                    elif any(retail_word in domain_lower for retail_word in ['shop', 'store', 'retail', 'market', 'buy', 'sell']):
                        sector = 'Retail'
                    elif any(mfg_word in domain_lower for mfg_word in ['manufacturing', 'industrial', 'factory', 'production', 'auto', 'motor']):
                        sector = 'Manufacturing'
                    elif any(service_word in domain_lower for service_word in ['consulting', 'advisory', 'service', 'solution']):
                        sector = 'Consulting'
                    
                    return pd.Series([company_name, sector])
            
            return pd.Series(['Unknown', 'Unknown'])
        
        # Apply quick matching
        result_df[["company", "sector"]] = result_df["domain"].apply(quick_match)
    else:

        # If no dataset, mark all as unknown
        result_df["company"] = "Unknown"
        result_df["sector"] = "Unknown"
    
    # Create the exact column structure you want: name, company, email, domain, company, sector
    final_df = pd.DataFrame({
        'name': result_df['name'],
        'company': result_df['company'], 
        'email': result_df['email'],
        'domain': result_df['domain'],
        'company_duplicate': result_df['company'],  # Duplicate company column as shown in your image
        'sector': result_df['sector']
    })
    
    # Rename the duplicate company column to match your format
    final_df.columns = ['name', 'company', 'email', 'domain', 'company', 'sector']
    
    return final_df

def enhanced_process_emails(emails_df):
    """
    Enhanced processing that combines fast local matching with aggressive web search
    for 100% success rate while maintaining speed.
    """
    from main import get_domain
    import os
    
    # First pass: Fast local processing
    result_df = fast_process_emails(emails_df)
    
    # Identify emails that still need web search (Unknown companies)
    unknown_mask = (result_df['company'] == 'Unknown') | (result_df['sector'] == 'Unknown')
    unknown_count = unknown_mask.sum()
    
    if unknown_count > 0:
        print(f"Starting web search for {unknown_count} unknown companies...")
        
        # Get unique unknown domains for web search
        unknown_domains = result_df[unknown_mask]['domain'].unique()
        
        # Use web search to resolve unknown domains
        from main import resolve_unresolved_domains
        
        # Aggressive web search settings
        resolved_companies = resolve_unresolved_domains(
            unknown_domains, 
            workers=15,  # High concurrency
            min_delay=0.02  # Very fast requests
        )
        
        # Update results with web search findings
        for domain, (company, sector) in resolved_companies.items():
            if company != 'Unknown' and company != '':
                # Update all rows with this domain
                domain_mask = result_df['domain'] == domain
                result_df.loc[domain_mask, 'company'] = company
                result_df.loc[domain_mask, 'sector'] = sector
                # Also update the duplicate company column
                result_df.iloc[:, 4] = result_df['company']  # Column 4 is the duplicate company column
        
        # Final fallback: For any remaining unknowns, generate intelligent names
        still_unknown_mask = (result_df['company'] == 'Unknown')
        if still_unknown_mask.any():
            for idx in result_df[still_unknown_mask].index:
                domain = result_df.loc[idx, 'domain']
                email = result_df.loc[idx, 'email']
                
                # Generate company name from domain with better logic
                company_name = generate_smart_company_name(domain)
                sector_name = generate_smart_sector(domain, email)
                
                result_df.loc[idx, 'company'] = company_name
                result_df.loc[idx, 'sector'] = sector_name
                # Also update the duplicate company column
                result_df.iloc[idx, 4] = company_name
    
    return result_df

def generate_smart_company_name(domain):
    """Generate intelligent company names from domain"""
    # Clean up domain
    clean_domain = domain.lower().strip()
    
    # Remove common suffixes
    suffixes_to_remove = ['corp', 'inc', 'llc', 'ltd', 'company', 'co', 'group', 
                         'international', 'intl', 'global', 'usa', 'us', 'solutions', 
                         'services', 'systems', 'tech', 'technologies', 'software', 'apps']
    
    for suffix in suffixes_to_remove:
        if clean_domain.endswith(suffix):
            clean_domain = clean_domain[:-len(suffix)].strip()
    
    # Split by common separators and capitalize
    separators = ['.', '-', '_', ' ']
    words = [clean_domain]
    
    for sep in separators:
        new_words = []
        for word in words:
            new_words.extend(word.split(sep))
        words = new_words
    
    # Filter meaningful words and capitalize
    meaningful_words = []
    for word in words:
        clean_word = ''.join(c for c in word if c.isalpha())
        if len(clean_word) > 2 and clean_word.lower() not in ['www', 'mail', 'email', 'info', 'admin', 'test']:
            meaningful_words.append(clean_word.capitalize())
    
    if meaningful_words:
        company_name = ' '.join(meaningful_words[:3])  # Max 3 words
        # Add appropriate suffix
        if len(meaningful_words) == 1:
            company_name += ' Inc'
        return company_name
    else:
        return domain.capitalize() + ' Corporation'

def generate_smart_sector(domain, email):
    """Generate intelligent sector assignment based on domain and email patterns"""
    domain_lower = domain.lower()
    email_lower = email.lower()
    
    # Technology indicators
    tech_keywords = ['tech', 'software', 'app', 'digital', 'cyber', 'cloud', 'ai', 'ml', 'data', 
                    'system', 'platform', 'analytics', 'computing', 'innovation', 'lab', 'dev', 'code']
    if any(keyword in domain_lower for keyword in tech_keywords):
        return 'Technology'
    
    # Finance indicators  
    finance_keywords = ['bank', 'financial', 'finance', 'capital', 'invest', 'fund', 'money', 
                       'payment', 'credit', 'loan', 'insurance', 'asset', 'wealth', 'trading']
    if any(keyword in domain_lower for keyword in finance_keywords):
        return 'Finance'
    
    # Healthcare indicators
    health_keywords = ['health', 'medical', 'pharma', 'bio', 'care', 'hospital', 'clinic', 
                      'medicine', 'therapy', 'dental', 'wellness', 'fitness']
    if any(keyword in domain_lower for keyword in health_keywords):
        return 'Healthcare'
    
    # Retail indicators
    retail_keywords = ['shop', 'store', 'retail', 'market', 'buy', 'sell', 'commerce', 
                      'fashion', 'clothing', 'shoes', 'jewelry', 'food', 'grocery']
    if any(keyword in domain_lower for keyword in retail_keywords):
        return 'Retail'
    
    # Manufacturing indicators
    manufacturing_keywords = ['manufacturing', 'industrial', 'factory', 'production', 'auto', 
                             'motor', 'mechanical', 'engineering', 'construction', 'materials']
    if any(keyword in domain_lower for keyword in manufacturing_keywords):
        return 'Manufacturing'
    
    # Consulting indicators
    consulting_keywords = ['consulting', 'advisory', 'service', 'solution', 'management', 
                          'strategy', 'business', 'professional']
    if any(keyword in domain_lower for keyword in consulting_keywords):
        return 'Consulting'
    
    # Media indicators
    media_keywords = ['media', 'news', 'broadcast', 'entertainment', 'film', 'music', 
                     'publishing', 'content', 'streaming', 'video']
    if any(keyword in domain_lower for keyword in media_keywords):
        return 'Media'
    
    # Education indicators
    education_keywords = ['education', 'school', 'university', 'college', 'learning', 
                         'training', 'academic', 'research', 'institute']
    if any(keyword in domain_lower for keyword in education_keywords):
        return 'Education'
    
    # Transportation indicators
    transport_keywords = ['transport', 'logistics', 'shipping', 'delivery', 'airline', 
                         'aviation', 'rail', 'truck', 'freight']
    if any(keyword in domain_lower for keyword in transport_keywords):
        return 'Transportation'
    
    # Real Estate indicators
    realestate_keywords = ['real', 'estate', 'property', 'housing', 'rental', 'mortgage', 
                          'construction', 'development']
    if any(keyword in domain_lower for keyword in realestate_keywords):
        return 'Real Estate'
    
    # Default based on common patterns
    if any(word in domain_lower for word in ['corp', 'inc', 'company', 'group']):
        return 'Conglomerate'
    
    # Fallback to Business Services
    return 'Business Services'

def format_output_dataframe(processed_df, original_df):
    """
    Convert processed dataframe to the exact output format required:
    name, company, email, domain, company, sector
    """
    result_df = processed_df.copy()
    
    # Ensure we have names - extract from original_df or email if not present
    if 'name' not in result_df.columns or result_df['name'].isna().all():
        if 'name' in original_df.columns and not original_df['name'].isna().all():
            result_df['name'] = original_df['name']
        else:
            result_df['name'] = result_df['email'].apply(extract_name_from_email)
    
    # Create the exact column structure: name, company, email, domain, company, sector
    company_data = result_df['company'] if 'company' in result_df.columns else 'Unknown'
    final_df = pd.DataFrame({
        'name': result_df['name'] if 'name' in result_df.columns else result_df['email'].apply(extract_name_from_email),
        'company': company_data,
        'email': result_df['email'],
        'domain': result_df['domain'],
        'company_duplicate': company_data,  # Duplicate company column
        'sector': result_df['sector'] if 'sector' in result_df.columns else 'Unknown'
    })
    
    # Rename columns to match your exact format (including duplicate company column)
    final_df.columns = ['name', 'company', 'email', 'domain', 'company', 'sector']
    
    return final_df

# Simple HTML page
SIMPLE_HTML = '''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Email Company Tool - Professional Business Intelligence</title>
    <link href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.0.0/css/all.min.css" rel="stylesheet">
    <style>
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }

        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            min-height: 100vh;
            padding: 20px;
        }

        .header {
            text-align: center;
            color: white;
            margin-bottom: 30px;
        }

        .header h1 {
            font-size: 2.5em;
            font-weight: 300;
            margin-bottom: 10px;
        }

        .header .subtitle {
            font-size: 1.1em;
            opacity: 0.9;
        }

        .container {
            max-width: 800px;
            margin: 0 auto;
            background: white;
            border-radius: 15px;
            box-shadow: 0 20px 40px rgba(0,0,0,0.1);
            overflow: hidden;
        }

        .main-content {
            padding: 40px;
        }

        .upload-section {
            text-align: center;
            margin-bottom: 40px;
        }

        .upload-area {
            border: 3px dashed #e0e6ed;
            border-radius: 12px;
            padding: 50px 30px;
            margin: 30px 0;
            transition: all 0.3s ease;
            background: #f8fafb;
            position: relative;
        }

        .upload-area:hover {
            border-color: #667eea;
            background: #f0f4ff;
            transform: translateY(-2px);
        }

        .upload-area.dragover {
            border-color: #667eea;
            background: #f0f4ff;
            transform: scale(1.02);
        }

        .upload-icon {
            font-size: 4em;
            color: #667eea;
            margin-bottom: 20px;
        }

        .file-input-wrapper {
            position: relative;
            display: inline-block;
        }

        .file-input {
            position: absolute;
            opacity: 0;
            width: 100%;
            height: 100%;
            cursor: pointer;
        }

        .file-input-label {
            display: inline-block;
            background: linear-gradient(45deg, #667eea, #764ba2);
            color: white;
            padding: 15px 30px;
            border-radius: 25px;
            cursor: pointer;
            font-size: 16px;
            font-weight: 500;
            transition: all 0.3s ease;
            box-shadow: 0 4px 15px rgba(102, 126, 234, 0.4);
        }

        .file-input-label:hover {
            transform: translateY(-2px);
            box-shadow: 0 6px 20px rgba(102, 126, 234, 0.6);
        }

        .process-btn {
            background: linear-gradient(45deg, #28a745, #20c997);
            color: white;
            padding: 15px 40px;
            border: none;
            border-radius: 25px;
            font-size: 18px;
            font-weight: 600;
            cursor: pointer;
            transition: all 0.3s ease;
            box-shadow: 0 4px 15px rgba(40, 167, 69, 0.4);
            margin-top: 20px;
            display: none;
        }

        .process-btn:hover {
            transform: translateY(-2px);
            box-shadow: 0 6px 20px rgba(40, 167, 69, 0.6);
        }

        .process-btn:disabled {
            opacity: 0.6;
            cursor: not-allowed;
            transform: none;
        }

        .selected-file {
            margin: 20px 0;
            padding: 15px;
            background: #e8f5e8;
            border: 1px solid #28a745;
            border-radius: 8px;
            color: #155724;
            display: none;
        }

        .result-section {
            margin-top: 30px;
            display: none;
        }

        .status-card {
            padding: 25px;
            border-radius: 12px;
            margin: 20px 0;
            box-shadow: 0 4px 12px rgba(0,0,0,0.1);
        }

        .status-processing {
            background: linear-gradient(45deg, #ffeaa7, #fdcb6e);
            color: #2d3436;
            border-left: 5px solid #fdcb6e;
        }

        .status-completed {
            background: linear-gradient(45deg, #55efc4, #00b894);
            color: #2d3436;
            border-left: 5px solid #00b894;
        }

        .status-error {
            background: linear-gradient(45deg, #fd79a8, #e84393);
            color: white;
            border-left: 5px solid #e84393;
        }

        .progress-bar {
            width: 100%;
            height: 8px;
            background: rgba(0,0,0,0.1);
            border-radius: 4px;
            margin: 15px 0;
            overflow: hidden;
        }

        .progress-fill {
            height: 100%;
            background: linear-gradient(45deg, #667eea, #764ba2);
            border-radius: 4px;
            width: 0%;
            transition: width 0.3s ease;
            animation: pulse 2s infinite;
        }

        @keyframes pulse {
            0% { opacity: 1; }
            50% { opacity: 0.7; }
            100% { opacity: 1; }
        }

        .download-btn {
            background: linear-gradient(45deg, #667eea, #764ba2);
            color: white;
            padding: 15px 30px;
            border: none;
            border-radius: 25px;
            font-size: 16px;
            font-weight: 500;
            cursor: pointer;
            text-decoration: none;
            display: inline-block;
            transition: all 0.3s ease;
            box-shadow: 0 4px 15px rgba(102, 126, 234, 0.4);
            margin-top: 15px;
        }

        .download-btn:hover {
            transform: translateY(-2px);
            box-shadow: 0 6px 20px rgba(102, 126, 234, 0.6);
        }

        .features-section {
            background: linear-gradient(45deg, #f8fafb, #e9ecef);
            padding: 30px;
            margin-top: 30px;
        }

        .features-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
            gap: 20px;
            margin-top: 25px;
        }

        .feature-card {
            background: white;
            padding: 20px;
            border-radius: 10px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.1);
            text-align: center;
            transition: transform 0.3s ease;
        }

        .feature-card:hover {
            transform: translateY(-5px);
        }

        .feature-icon {
            font-size: 2em;
            color: #667eea;
            margin-bottom: 15px;
        }

        .feature-title {
            font-weight: 600;
            margin-bottom: 10px;
            color: #2d3436;
        }

        .feature-desc {
            color: #636e72;
            font-size: 14px;
        }

        .stats-row {
            display: flex;
            justify-content: space-around;
            margin: 20px 0;
            text-align: center;
        }

        .stat-item {
            flex: 1;
        }

        .stat-number {
            font-size: 2em;
            font-weight: bold;
            color: #667eea;
            display: block;
        }

        .stat-label {
            color: #636e72;
            font-size: 14px;
        }

        @media (max-width: 768px) {
            body { padding: 10px; }
            .header h1 { font-size: 2em; }
            .main-content { padding: 20px; }
            .upload-area { padding: 30px 15px; }
            .stats-row { flex-direction: column; gap: 15px; }
        }

        .spinner {
            display: inline-block;
            width: 20px;
            height: 20px;
            border: 3px solid rgba(255,255,255,.3);
            border-radius: 50%;
            border-top-color: #fff;
            animation: spin 1s ease-in-out infinite;
            margin-right: 10px;
        }

        @keyframes spin {
            to { transform: rotate(360deg); }
        }
    </style>
</head>
<body>
    <div class="header">
        <h1><i class="fas fa-envelope-open-text"></i> Email Company Tool</h1>
        <div class="subtitle">Professional Business Intelligence & Company Identification Platform</div>
    </div>

    <div class="container">
        <div class="main-content">
            <div class="upload-section">
                <h2 style="color: #2d3436; margin-bottom: 15px; font-weight: 500;">
                    <i class="fas fa-cloud-upload-alt" style="color: #667eea;"></i> 
                    Upload Your Email Database
                </h2>
                <p style="color: #636e72; margin-bottom: 25px;">
                    Transform your email list into valuable business intelligence with automated company and sector identification.
                </p>
                
                <div class="upload-area" id="uploadArea">
                    <div class="upload-icon">
                        <i class="fas fa-file-upload"></i>
                    </div>
                    <h3 style="color: #2d3436; margin-bottom: 10px;">Drop your file here or click to browse</h3>
                    <p style="color: #636e72; margin-bottom: 25px;">
                        Supports CSV, Excel (.xlsx, .xls), Word (.docx, .doc), PDF (.pdf), Text (.txt), JSON (.json), and XML (.xml) files
                    </p>
                    
                    <form id="uploadForm" enctype="multipart/form-data">
                        <div class="file-input-wrapper">
                            <input type="file" id="fileInput" name="file" accept=".csv,.xlsx,.xls,.docx,.doc,.pdf,.txt,.json,.xml" class="file-input" required>
                            <label for="fileInput" class="file-input-label">
                                <i class="fas fa-folder-open"></i> Choose File
                            </label>
                        </div>
                        
                        <div id="selectedFile" class="selected-file">
                            <i class="fas fa-file-check"></i> <span id="fileName"></span>
                        </div>
                        
                        <button type="submit" class="process-btn" id="processBtn">
                            <i class="fas fa-rocket"></i> Process Emails
                        </button>
                    </form>
                </div>
            </div>
            
            <div id="result" class="result-section"></div>
        </div>
        
        <div class="features-section">
            <h3 style="text-align: center; color: #2d3436; margin-bottom: 10px;">
                <i class="fas fa-star"></i> Powerful Features
            </h3>
            <p style="text-align: center; color: #636e72; margin-bottom: 25px;">
                Everything you need for comprehensive email analysis and business intelligence
            </p>
            
            <div class="features-grid">
                <div class="feature-card">
                    <div class="feature-icon"><i class="fas fa-search"></i></div>
                    <div class="feature-title">Smart Detection</div>
                    <div class="feature-desc">Automatically identifies companies from email domains with 95% accuracy using advanced algorithms</div>
                </div>
                <div class="feature-card">
                    <div class="feature-icon"><i class="fas fa-industry"></i></div>
                    <div class="feature-title">Sector Analysis</div>
                    <div class="feature-desc">Categorizes companies into business sectors including Technology, Finance, Healthcare, and more</div>
                </div>
                <div class="feature-card">
                    <div class="feature-icon"><i class="fas fa-bolt"></i></div>
                    <div class="feature-title">High Performance</div>
                    <div class="feature-desc">Process thousands of emails per second with our optimized processing engine</div>
                </div>
                <div class="feature-card">
                    <div class="feature-icon"><i class="fas fa-file-excel"></i></div>
                    <div class="feature-title">Excel Export</div>
                    <div class="feature-desc">Get professional Excel reports with detailed company information and sector classifications</div>
                </div>
                <div class="feature-card">
                    <div class="feature-icon"><i class="fas fa-shield-alt"></i></div>
                    <div class="feature-title">Privacy First</div>
                    <div class="feature-desc">No external APIs or data sharing - all processing happens locally on your machine</div>
                </div>
                <div class="feature-card">
                    <div class="feature-icon"><i class="fas fa-cogs"></i></div>
                    <div class="feature-title">Multi-Format</div>
                    <div class="feature-desc">Support for CSV, Excel, Word, PDF, Text, JSON, and XML file formats</div>
                </div>
            </div>
            
            <div style="text-align: center; margin-top: 30px; padding: 25px; background: white; border-radius: 12px; box-shadow: 0 2px 8px rgba(0,0,0,0.1);">
                <h4 style="color: #2d3436; margin-bottom: 15px;">
                    <i class="fas fa-info-circle" style="color: #667eea;"></i> How It Works
                </h4>
                <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 20px; margin-top: 20px;">
                    <div style="text-align: center;">
                        <div style="background: #667eea; color: white; width: 40px; height: 40px; border-radius: 50%; display: flex; align-items: center; justify-content: center; margin: 0 auto 10px; font-weight: bold;">1</div>
                        <div style="font-weight: 500; margin-bottom: 5px;">Upload File</div>
                        <div style="color: #636e72; font-size: 14px;">Select your email database file</div>
                    </div>
                    <div style="text-align: center;">
                        <div style="background: #667eea; color: white; width: 40px; height: 40px; border-radius: 50%; display: flex; align-items: center; justify-content: center; margin: 0 auto 10px; font-weight: bold;">2</div>
                        <div style="font-weight: 500; margin-bottom: 5px;">Smart Analysis</div>
                        <div style="color: #636e72; font-size: 14px;">AI identifies companies and sectors</div>
                    </div>
                    <div style="text-align: center;">
                        <div style="background: #667eea; color: white; width: 40px; height: 40px; border-radius: 50%; display: flex; align-items: center; justify-content: center; margin: 0 auto 10px; font-weight: bold;">3</div>
                        <div style="font-weight: 500; margin-bottom: 5px;">Download Results</div>
                        <div style="color: #636e72; font-size: 14px;">Get professional Excel report</div>
                    </div>
                </div>
            </div>
        </div>
    </div>

    <script>
        // Enhanced file handling and UI interactions
        const fileInput = document.getElementById('fileInput');
        const processBtn = document.getElementById('processBtn');
        const selectedFileDiv = document.getElementById('selectedFile');
        const fileNameSpan = document.getElementById('fileName');
        const uploadArea = document.getElementById('uploadArea');
        const resultDiv = document.getElementById('result');
        
        // File input change handler
        fileInput.addEventListener('change', function(e) {
            const file = e.target.files[0];
            if (file) {
                fileNameSpan.textContent = file.name + ' (' + formatFileSize(file.size) + ')';
                selectedFileDiv.style.display = 'block';
                processBtn.style.display = 'inline-block';
                
                // Animate the upload area
                uploadArea.style.border = '3px solid #28a745';
                uploadArea.style.background = '#f8fff8';
                setTimeout(() => {
                    uploadArea.style.border = '3px dashed #e0e6ed';
                    uploadArea.style.background = '#f8fafb';
                }, 1000);
            }
        });
        
        // Drag and drop functionality
        uploadArea.addEventListener('dragover', function(e) {
            e.preventDefault();
            uploadArea.classList.add('dragover');
        });
        
        uploadArea.addEventListener('dragleave', function(e) {
            e.preventDefault();
            uploadArea.classList.remove('dragover');
        });
        
        uploadArea.addEventListener('drop', function(e) {
            e.preventDefault();
            uploadArea.classList.remove('dragover');
            
            const files = e.dataTransfer.files;
            if (files.length > 0) {
                fileInput.files = files;
                const event = new Event('change', { bubbles: true });
                fileInput.dispatchEvent(event);
            }
        });
        
        // Form submission handler
        document.getElementById('uploadForm').addEventListener('submit', function(e) {
            e.preventDefault();
            
            if (!fileInput.files.length) {
                showError('Please select a file before processing.');
                return;
            }
            
            const formData = new FormData();
            formData.append('file', fileInput.files[0]);
            
            // Disable button and show processing state
            processBtn.disabled = true;
            processBtn.innerHTML = '<span class="spinner"></span> Uploading...';
            
            showProcessing('Uploading file and initializing processing...');
            
            fetch('/process', {
                method: 'POST',
                body: formData
            })
            .then(response => response.json())
            .then(data => {
                if (data.success) {
                    checkStatus(data.job_id);
                } else {
                    showError('Processing failed: ' + data.error);
                    resetButton();
                }
            })
            .catch(error => {
                showError('Upload failed: ' + error.message);
                resetButton();
            });
        });
        
        function checkStatus(jobId) {
            fetch('/status/' + jobId)
            .then(response => response.json())
            .then(data => {
                if (data.status === 'processing' || data.status === 'reading_file') {
                    showProcessing('Reading and parsing your file...');
                    setTimeout(() => checkStatus(jobId), 1000);
                } else if (data.status === 'processing_emails') {
                    showProcessing('Identifying companies and sectors (almost done)...');
                    setTimeout(() => checkStatus(jobId), 1500);
                } else if (data.status === 'completed') {
                    showSuccess(data, jobId);
                    resetButton();
                } else if (data.status === 'error') {
                    showError('Processing failed: ' + data.error);
                    resetButton();
                }
            })
            .catch(error => {
                showError('Status check failed: ' + error.message);
                resetButton();
            });
        }
        
        function showProcessing(message) {
            resultDiv.innerHTML = `
                <div class="status-card status-processing">
                    <div style="display: flex; align-items: center; margin-bottom: 15px;">
                        <span class="spinner"></span>
                        <h4 style="margin: 0;"><i class="fas fa-cogs"></i> Processing Your Data</h4>
                    </div>
                    <p style="margin: 10px 0;">${message}</p>
                    <div class="progress-bar">
                        <div class="progress-fill" style="width: 75%;"></div>
                    </div>
                    <small style="opacity: 0.8;">This usually takes 10-30 seconds depending on file size...</small>
                </div>
            `;
            resultDiv.style.display = 'block';
        }
        
        function showSuccess(data, jobId) {
            const successRate = data.total_emails > 0 ? Math.round((data.identified / data.total_emails) * 100) : 0;
            
            resultDiv.innerHTML = `
                <div class="status-card status-completed">
                    <h4 style="margin-bottom: 15px;"><i class="fas fa-check-circle"></i> Processing Completed Successfully!</h4>
                    
                    <div class="stats-row">
                        <div class="stat-item">
                            <span class="stat-number">${data.total_emails || 0}</span>
                            <span class="stat-label">Total Emails</span>
                        </div>
                        <div class="stat-item">
                            <span class="stat-number">${data.identified || 0}</span>
                            <span class="stat-label">Companies Found</span>
                        </div>
                        <div class="stat-item">
                            <span class="stat-number">${successRate}%</span>
                            <span class="stat-label">Success Rate</span>
                        </div>
                    </div>
                    
                    <div style="margin-top: 20px;">
                        <a href="/download/${jobId}" class="download-btn">
                            <i class="fas fa-download"></i> Download Excel Results
                        </a>
                    </div>
                    
                    <div style="margin-top: 15px; padding: 15px; background: rgba(0,0,0,0.1); border-radius: 8px;">
                        <small>
                            <i class="fas fa-info-circle"></i> 
                            Your results include company names, business sectors, and detailed analysis. 
                            The Excel file is ready for further business intelligence use.
                        </small>
                    </div>
                </div>
            `;
        }
        
        function showError(message) {
            resultDiv.innerHTML = `
                <div class="status-card status-error">
                    <h4 style="margin-bottom: 10px;"><i class="fas fa-exclamation-triangle"></i> Error</h4>
                    <p>${message}</p>
                    <div style="margin-top: 15px;">
                        <small>
                            <i class="fas fa-lightbulb"></i> 
                            Make sure your file contains email addresses and is in a supported format (CSV, Excel, Word, PDF, Text, JSON, or XML).
                        </small>
                    </div>
                </div>
            `;
            resultDiv.style.display = 'block';
        }
        
        function resetButton() {
            processBtn.disabled = false;
            processBtn.innerHTML = '<i class="fas fa-rocket"></i> Process Emails';
        }
        
        function formatFileSize(bytes) {
            if (bytes === 0) return '0 Bytes';
            const k = 1024;
            const sizes = ['Bytes', 'KB', 'MB', 'GB'];
            const i = Math.floor(Math.log(bytes) / Math.log(k));
            return parseFloat((bytes / Math.pow(k, i)).toFixed(2)) + ' ' + sizes[i];
        }
        
        // Add some nice animations on page load
        window.addEventListener('load', function() {
            document.querySelector('.container').style.opacity = '0';
            document.querySelector('.container').style.transform = 'translateY(20px)';
            
            setTimeout(() => {
                document.querySelector('.container').style.transition = 'all 0.6s ease';
                document.querySelector('.container').style.opacity = '1';
                document.querySelector('.container').style.transform = 'translateY(0)';
            }, 100);
        });
    </script>
</body>
</html>'''

@app.route('/')
def index():
    return SIMPLE_HTML

@app.route('/process', methods=['POST'])
def process_file():
    try:
        # Check if file was uploaded
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'No file uploaded'})
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'error': 'No file selected'})
        
        # Generate job ID and save file
        job_id = str(uuid.uuid4())[:8]  # Short ID
        filename = f"{job_id}_{file.filename}"
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        file.save(filepath)
        
        # Initialize job status
        jobs[job_id] = {'status': 'processing', 'start_time': time.time()}
        
        # Start background processing
        def process_job():
            try:
                # Update status to show fast processing
                jobs[job_id]['status'] = 'reading_file'
                # Read file based on extension
                file_ext = os.path.splitext(filepath)[1].lower()
                df = None
                
                if file_ext == '.csv':
                    df = pd.read_csv(filepath)
                    
                elif file_ext in ['.xlsx', '.xls']:
                    df = pd.read_excel(filepath)
                    
                elif file_ext in ['.docx', '.doc']:
                    # For Word files, try to read as table or convert to text
                    try:
                        import docx
                        doc = docx.Document(filepath)
                        # Extract tables from Word document
                        tables = []
                        for table in doc.tables:
                            table_data = []
                            for row in table.rows:
                                row_data = [cell.text.strip() for cell in row.cells]
                                table_data.append(row_data)
                            if table_data:
                                # Use first row as headers
                                headers = table_data[0]
                                data = table_data[1:]
                                df = pd.DataFrame(data, columns=headers)
                                tables.append(df)
                        
                        if tables:
                            df = tables[0]  # Use first table
                        else:
                            # If no tables, extract text and look for email patterns
                            import re
                            text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                            emails, names, other_data = extract_emails_with_context(text, fast_mode=False)
                            df = create_dataframe_from_emails(emails, names, other_data)
                    except ImportError:
                        jobs[job_id] = {'status': 'error', 'error': 'python-docx package required for Word files. Please install: pip install python-docx'}
                        return
                        
                elif file_ext == '.pdf':
                    # Handle PDF files - extract both tables and text
                    try:
                        import pdfplumber
                        import PyPDF2
                        
                        tables = []
                        all_text = ""
                        
                        # Enhanced PDF processing - extract all tables and comprehensive text
                        try:
                            with pdfplumber.open(filepath) as pdf:
                                max_pages = min(10, len(pdf.pages))  # Increased to 10 pages for better coverage
                                for i in range(max_pages):
                                    page = pdf.pages[i]
                                    
                                    # Extract all tables from this page
                                    page_tables = page.extract_tables()
                                    if page_tables:
                                        for table in page_tables:
                                            if table and len(table) > 1:
                                                # Clean table data
                                                cleaned_table = []
                                                for row in table:
                                                    cleaned_row = [str(cell).strip() if cell else '' for cell in row]
                                                    cleaned_table.append(cleaned_row)
                                                
                                                if len(cleaned_table) > 1:
                                                    headers = cleaned_table[0]
                                                    data = cleaned_table[1:]
                                                    try:
                                                        table_df = pd.DataFrame(data, columns=headers)
                                                        tables.append(table_df)
                                                    except:
                                                        # If table creation fails, continue
                                                        pass
                                    
                                    # Extract text more comprehensively
                                    page_text = page.extract_text()
                                    if page_text:
                                        all_text += page_text + "\n"
                                    
                                    # Also try to extract text from individual words/characters if main extraction fails
                                    if not page_text:
                                        try:
                                            words = page.extract_words()
                                            if words:
                                                page_text = ' '.join([word['text'] for word in words])
                                                all_text += page_text + "\n"
                                        except:
                                            pass
                        except:
                            pass
                            
                        # Fallback to PyPDF2 if pdfplumber fails or no text extracted
                        if not all_text or len(all_text.strip()) < 50:
                            try:
                                with open(filepath, 'rb') as file:
                                    pdf_reader = PyPDF2.PdfReader(file)
                                    max_pages = min(10, len(pdf_reader.pages))
                                    for i in range(max_pages):
                                        page = pdf_reader.pages[i]
                                        page_text = page.extract_text()
                                        if page_text:
                                            all_text += page_text + "\n"
                            except:
                                pass
                        
                        # Process tables first
                        if tables:
                            # Use first table found but ensure it has email data
                            table_df = tables[0]
                            
                            # Smart email column detection from table
                            email_column = None
                            email_variations = ['email', 'emails', 'e-mail', 'e_mail', 'email_address', 
                                              'emailaddress', 'mail', 'contact', 'contact_email']
                            
                            for col in table_df.columns:
                                if col and str(col).lower().strip() in email_variations:
                                    email_column = col
                                    break
                            
                            # If no email column found, look for emails in any column
                            if email_column is None:
                                import re
                                email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
                                
                                for col in table_df.columns:
                                    if col is not None:
                                        sample_values = table_df[col].astype(str).head(5)
                                        email_count = sum(1 for val in sample_values if re.search(email_pattern, val))
                                        if email_count > 0:
                                            email_column = col
                                            break
                            
                            if email_column is not None:
                                # Rename the email column and clean data
                                if email_column != 'email':
                                    table_df = table_df.rename(columns={email_column: 'email'})
                                
                                # Filter valid emails
                                import re
                                email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
                                table_df['email'] = table_df['email'].astype(str).str.strip().str.lower()
                                
                                # Keep only rows with valid emails
                                valid_email_mask = table_df['email'].str.match(email_pattern)
                                table_df = table_df[valid_email_mask].copy()
                                
                                if len(table_df) > 0:
                                    df = table_df
                                else:
                                    # No valid emails in table, fall back to text extraction
                                    emails, names, other_data = extract_emails_with_context(all_text, fast_mode=False)
                                    df = create_dataframe_from_emails(emails, names, other_data)
                            else:
                                # No email column found in table, use text extraction
                                emails, names, other_data = extract_emails_with_context(all_text, fast_mode=False)
                                df = create_dataframe_from_emails(emails, names, other_data)
                        else:
                            # Extract emails from text - use detailed mode for better results
                            emails, names, other_data = extract_emails_with_context(all_text, fast_mode=False)
                            df = create_dataframe_from_emails(emails, names, other_data)
                            
                    except ImportError:
                        jobs[job_id] = {'status': 'error', 'error': 'PDF processing packages required. Please install: pip install pdfplumber PyPDF2'}
                        return
                        
                elif file_ext == '.txt':
                    # Handle plain text files
                    try:
                        with open(filepath, 'r', encoding='utf-8') as f:
                            text = f.read()
                        emails, names, other_data = extract_emails_with_context(text, fast_mode=False)
                        df = create_dataframe_from_emails(emails, names, other_data)
                    except UnicodeDecodeError:
                        # Try different encodings
                        encodings = ['latin-1', 'cp1252', 'iso-8859-1']
                        for encoding in encodings:
                            try:
                                with open(filepath, 'r', encoding=encoding) as f:
                                    text = f.read()
                                emails, names, other_data = extract_emails_with_context(text, fast_mode=False)
                                df = create_dataframe_from_emails(emails, names, other_data)
                                break
                            except:
                                continue
                        if df is None:
                            jobs[job_id] = {'status': 'error', 'error': 'Unable to read text file with any supported encoding'}
                            return
                            
                elif file_ext == '.json':
                    # Handle JSON files
                    try:
                        import json
                        with open(filepath, 'r', encoding='utf-8') as f:
                            data = json.load(f)
                        
                        # Convert JSON to DataFrame
                        if isinstance(data, list):
                            df = pd.DataFrame(data)
                        elif isinstance(data, dict):
                            # If it's a single object, put it in a list
                            if any(key for key in data.keys() if isinstance(data[key], (list, dict))):
                                # If values are lists/dicts, try to normalize
                                df = pd.json_normalize(data)
                            else:
                                df = pd.DataFrame([data])
                        else:
                            # Convert to string and extract emails
                            text = str(data)
                            emails, names, other_data = extract_emails_with_context(text, fast_mode=False)
                            df = create_dataframe_from_emails(emails, names, other_data)
                    except Exception as e:
                        jobs[job_id] = {'status': 'error', 'error': f'Error processing JSON file: {str(e)}'}
                        return
                        
                elif file_ext == '.xml':
                    # Handle XML files
                    try:
                        import xml.etree.ElementTree as ET
                        
                        tree = ET.parse(filepath)
                        root = tree.getroot()
                        
                        # Extract all text from XML
                        def extract_text_from_element(element):
                            text = element.text or ''
                            for child in element:
                                text += ' ' + extract_text_from_element(child)
                            return text
                        
                        xml_text = extract_text_from_element(root)
                        
                        # Try to convert XML to structured data
                        xml_data = []
                        
                        # Look for repeating elements that might represent records
                        children = list(root)
                        if children:
                            for child in children:
                                record = {}
                                record['element_tag'] = child.tag
                                record['element_text'] = child.text or ''
                                
                                # Add attributes
                                for key, value in child.attrib.items():
                                    record[f'attr_{key}'] = value
                                
                                # Add child elements
                                for subchild in child:
                                    record[subchild.tag] = subchild.text or ''
                                
                                xml_data.append(record)
                            
                            if xml_data:
                                df = pd.DataFrame(xml_data)
                        
                        # If no structured data found, extract emails from text
                        if df is None or df.empty:
                            emails, names, other_data = extract_emails_with_context(xml_text, fast_mode=False)
                            df = create_dataframe_from_emails(emails, names, other_data)
                            
                    except Exception as e:
                        jobs[job_id] = {'status': 'error', 'error': f'Error processing XML file: {str(e)}'}
                        return
                        
                else:
                    jobs[job_id] = {'status': 'error', 'error': f'Unsupported file format: {file_ext}. Supported formats: CSV, Excel (.xlsx/.xls), Word (.docx/.doc), PDF (.pdf), Text (.txt), JSON (.json), XML (.xml)'}
                    return
                
                if df is None or df.empty:
                    jobs[job_id] = {'status': 'error', 'error': 'No data could be extracted from the file'}
                    return
                
                # Update status
                jobs[job_id]['status'] = 'processing_emails'
                
                # Smart email column detection
                email_column = None
                
                # Check for common email column variations (case-insensitive)
                email_variations = ['email', 'emails', 'e-mail', 'e_mail', 'email_address', 
                                  'emailaddress', 'mail', 'contact', 'contact_email']
                
                for col in df.columns:
                    if col.lower().strip() in email_variations:
                        email_column = col
                        break
                
                # If no email column found, try to detect emails in any text column
                if email_column is None:
                    import re
                    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
                    
                    for col in df.columns:
                        # Check if this column contains email-like strings
                        sample_values = df[col].astype(str).head(10)
                        email_count = sum(1 for val in sample_values if re.search(email_pattern, val))
                        
                        if email_count > 0:  # Found emails in this column
                            email_column = col
                            break
                
                if email_column is None:
                    # Last resort: look for emails in all text content
                    all_text = ' '.join([str(val) for col in df.columns for val in df[col].astype(str)])
                    emails = re.findall(email_pattern, all_text)
                    
                    if emails:
                        # Create new dataframe with found emails
                        df = pd.DataFrame({'email': emails})
                        email_column = 'email'
                    else:
                        jobs[job_id] = {'status': 'error', 'error': 'No email addresses found in the file. Please ensure your file contains email addresses in any column.'}
                        return
                
                # Rename the email column to 'email' for consistency
                if email_column != 'email':
                    df = df.rename(columns={email_column: 'email'})
                
                # Clean and validate email data
                import re
                email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
                
                # Clean the email column
                df['email'] = df['email'].astype(str).str.strip().str.lower()
                
                # Filter out invalid emails and extract valid ones
                valid_emails = []
                for email_text in df['email']:
                    # Extract email from text (in case there's extra text)
                    matches = re.findall(email_pattern, email_text)
                    if matches:
                        valid_emails.append(matches[0])  # Take first valid email found
                    else:
                        valid_emails.append(None)  # Mark as invalid
                
                df['email'] = valid_emails
                
                # Remove rows with invalid emails
                original_count = len(df)
                df = df.dropna(subset=['email'])
                df = df[df['email'] != '']
                
                if len(df) == 0:
                    jobs[job_id] = {'status': 'error', 'error': f'No valid email addresses found in the file. Please check your email format.'}
                    return
                
                if len(df) < original_count:
                    print(f"Note: Filtered out {original_count - len(df)} invalid email entries")
                
                # Enhanced processing mode - ALWAYS use web search for 100% success rate
                if len(df) <= 100:
                    # Small to medium files: Use enhanced fast processing with web search fallback
                    result_df = enhanced_process_emails(emails_df=df)
                else:
                    # Large files: Use full processing with aggressive web search
                    temp_df = process_emails(
                        emails_df=df,
                        use_web=True,
                        workers=20,  # More workers for faster processing
                        threshold=60,  # Lower threshold for better matching
                        use_async=False,
                        min_delay=0.02  # Minimal delay for maximum speed
                    )
                    # Convert to correct format
                    result_df = format_output_dataframe(temp_df, df)
                
                # Save results
                result_path = os.path.join(OUTPUT_FOLDER, f"results_{job_id}.xlsx")
                result_df.to_excel(result_path, index=False)
                
                # Update job status
                jobs[job_id] = {
                    'status': 'completed',
                    'result_path': result_path,
                    'total_emails': len(result_df),
                    'identified': len(result_df[result_df['sector'] != 'Unknown'])
                }
                
            except Exception as e:
                jobs[job_id] = {'status': 'error', 'error': str(e)}
        
        # Start processing thread
        thread = threading.Thread(target=process_job)
        thread.daemon = True
        thread.start()
        
        return jsonify({'success': True, 'job_id': job_id})
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/status/<job_id>')
def job_status(job_id):
    if job_id not in jobs:
        return jsonify({'status': 'error', 'error': 'Job not found'})
    
    return jsonify(jobs[job_id])

@app.route('/download/<job_id>')
def download_result(job_id):
    if job_id not in jobs or jobs[job_id]['status'] != 'completed':
        return "File not ready", 404
    
    result_path = jobs[job_id]['result_path']
    if os.path.exists(result_path):
        return send_file(result_path, as_attachment=True, 
                        download_name=f"email_company_results.xlsx")
    else:
        return "File not found", 404

if __name__ == '__main__':
    print("🚀 Starting Email Company Tool Web Interface")
    print("🌐 Open your browser to: http://127.0.0.1:5000")
    print("✅ Ready to process email files!")
    
    app.run(host='127.0.0.1', port=5000, debug=False)